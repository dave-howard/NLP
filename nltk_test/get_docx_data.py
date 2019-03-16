from docx import Document
import os
import io
basedir = os.path.abspath(os.path.dirname(__file__))


# get data from a single docx file
# return list[] of paragraph text entries
def read_docx(filename) -> list:
    file_name = os.path.join(basedir, "..", filename)
    print(file_name)
    with open(file_name, 'rb') as f:
        source_stream = io.BytesIO(f.read())
    document = Document(source_stream)
    source_stream.close()
    text = []
    for p in document.paragraphs:
        print(p.text)
        text.append(p.text+" ")
    return text


# get data from a folder of docx files
# return list[] of paragraph text entries
def read_docx_folder(folder) -> list:
    path = os.path.join(basedir, "..", folder)
    print(path)
    text = []
    for filename in os.listdir(path):
        file_path = os.path.join(basedir, "..", folder, filename)
        print(file_path)
        try:
            with open(file_path, 'rb') as f:
                source_stream = io.BytesIO(f.read())
            document = Document(source_stream)
            source_stream.close()
            doc_text = "["+filename+"]"
            for p in document.paragraphs:
                if len(p.text)> 50:
                    print(p.text)
                    doc_text+=p.text.replace("\n"," ")+" "
            text.append(doc_text)
        except:
            pass
    return text